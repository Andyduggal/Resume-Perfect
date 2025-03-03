import os
import streamlit as st
import groq
import docx
from dotenv import load_dotenv
from streamlit_ace import st_ace
from difflib import HtmlDiff
import io

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Initialize Groq Client
client = groq.Groq(api_key=GROQ_API_KEY)

def extract_text_from_docx(uploaded_file):
    """Extracts text from a Word document."""
    doc = docx.Document(uploaded_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def generate_updated_resume(resume_text, job_description):
    """Uses Groq LLM to enhance the resume to match job description as closely as possible."""
    prompt = f"""
    Given the following resume, improve it to match the job description as closely as possible while keeping the format intact.
    Ensure that all relevant skills, experience, and qualifications from the job description are reflected in the resume.
    
    **Job Description:**
    {job_description}
    
    **Original Resume:**
    {resume_text}
    
    **Updated Resume:**
    """

    response = client.chat.completions.create(
        model="llama3-8b-8192",
        messages=[
            {"role": "system", "content": "You are an expert resume writer. Ensure the resume aligns closely with the job description."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500,
        temperature=0.7
    )
    
    return response.choices[0].message.content.strip()

def generate_cover_letter(job_description, resume_text, tone="Professional"):
    """Generates a tailored cover letter based on job description and resume."""
    prompt = f"""
    Generate a {tone} and compelling cover letter based on the job description and resume provided below.
    
    **Job Description:**
    {job_description}
    
    **Resume:**
    {resume_text}
    
    **Cover Letter:**
    """
    
    response = client.chat.completions.create(
        model="llama3-8b-8192",
        messages=[
            {"role": "system", "content": "You are an expert cover letter writer. Ensure the cover letter is well-structured and personalized."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000,
        temperature=0.7
    )
    
    return response.choices[0].message.content.strip()

def calculate_match_percentage(resume_text, job_description):
    """Calculates the percentage match of resume to job description."""
    resume_words = set(resume_text.lower().split())
    job_words = set(job_description.lower().split())
    match = len(resume_words & job_words) / len(job_words) * 100 if job_words else 0
    return round(match, 2)

def create_word_doc(text):
    """Creates a Word document from the given text and returns it as a binary stream."""
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    binary_stream = io.BytesIO()
    doc.save(binary_stream)
    binary_stream.seek(0)
    return binary_stream

# Streamlit UI
st.title("🚀 AI-Powered Resume & Cover Letter Enhancer")
st.markdown("### Optimize your resume and generate a professional cover letter!")

uploaded_file = st.file_uploader("📂 Upload your resume (Word document)", type=["docx"])
job_description = st.text_area("📝 Paste the job description here", placeholder="Enter job details here...")

tone = st.selectbox("✍️ Cover Letter Tone", ["Professional", "Casual", "Creative", "Industry-Specific"], index=0)

if st.button("🔍 Start Process", disabled=not (uploaded_file and job_description)):
    with st.spinner("Processing your resume and cover letter... ⏳"):
        resume_text = extract_text_from_docx(uploaded_file)
        pre_update_match = calculate_match_percentage(resume_text, job_description)
        updated_resume_text = generate_updated_resume(resume_text, job_description)
        post_update_match = calculate_match_percentage(updated_resume_text, job_description)
        cover_letter_text = generate_cover_letter(job_description, updated_resume_text, tone)
        
        # Store the generated texts in session state
        st.session_state['updated_resume_text'] = updated_resume_text
        st.session_state['cover_letter_text'] = cover_letter_text
        st.session_state['pre_update_match'] = pre_update_match
        st.session_state['post_update_match'] = post_update_match
    
    st.success("✅ Process Completed!")

# Retrieve the generated texts from session state
updated_resume_text = st.session_state.get('updated_resume_text', '')
cover_letter_text = st.session_state.get('cover_letter_text', '')
pre_update_match = st.session_state.get('pre_update_match', 0)
post_update_match = st.session_state.get('post_update_match', 0)

if updated_resume_text and cover_letter_text:
    col1, col2 = st.columns(2)
    with col1:
        st.metric(label="Before Update", value=f"{pre_update_match}%")
    with col2:
        st.metric(label="After Update", value=f"{post_update_match}%")
    
    with st.expander("📜 Updated Resume (Editable)"):
        updated_resume_text = st_ace(value=updated_resume_text, language='markdown', theme='monokai', height=400)
    
    with st.expander("📜 Generated Cover Letter"):
        st.text_area("", cover_letter_text, height=300)
    
    st.subheader("📥 Download Your Documents")
    col1, col2 = st.columns(2)
    with col1:
        st.download_button("📥 Download Updated Resume", create_word_doc(updated_resume_text), file_name="Updated_Resume.docx")
    with col2:
        st.download_button("📥 Download Cover Letter", create_word_doc(cover_letter_text), file_name="Cover_Letter.docx")